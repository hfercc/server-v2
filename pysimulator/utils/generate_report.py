import docx
from docx import Document
from docx.shared import Inches
import os
from optparse import OptionParser

if __name__ == '__main__':
    parser = OptionParser()
    parser.add_option('--folder',dest='folder')
    (options, args) = parser.parse_args()
    
    alpha_id = os.path.basename(options.folder)
    category = ['longshort_ALL', 'longshort_zz500', 'longshort_hs300', 'IC_hedge', 'IF_hedge', 'longonly_ALL', 'longonly_zz500', 'longonly_hs300']
    
    document = Document()
    document.add_heading('Alpha {0} Report'.format(alpha_id), 0)

    formula_path = os.path.join(options.folder,'formula.txt')
    formula_f = open(formula_path, 'r')
    formula_str = ''
    for line in formula_f:
        if len(line) > 0:
            formula_str = line.strip('\n').strip('\r')

    p = document.add_paragraph('')
    p.add_run('Alpha Formula:{0}'.format(formula_str)).bold = True

    for ctgry in category:
        label = '{0}_{1}'.format(alpha_id, ctgry)        
        document.add_paragraph('{0}'.format(ctgry), style='List Bullet')
        
        table = document.add_table(rows=1, cols=8)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Period'
        hdr_cells[1].text = '%Tvr'
        hdr_cells[2].text = '%Ret'
        hdr_cells[3].text = '%CumRet'
        #hdr_cells[4].text = 'Pnl'
        #hdr_cells[5].text = 'CumPnl'
        hdr_cells[4].text = 'Sharpe'
        hdr_cells[5].text = 'IR'

        prf_path = os.path.join(options.folder, label + '_performance.csv')
        prf_f = open(prf_path, 'r')
        lc = 0
        for line in prf_f:
            lc += 1
            if lc == 1:
                continue
            row_cells = table.add_row().cells
            tokens = line.split()            
            row_cells[0].text = tokens[0]
            row_cells[1].text = tokens[1]
            row_cells[2].text = tokens[2]
            row_cells[3].text = tokens[3]
            row_cells[4].text = tokens[6]
            row_cells[5].text = tokens[7]
        
        fig_path = os.path.join(options.folder, label + '_pnl.png')
        document.add_picture(fig_path, width=Inches(5.5))

    document.save('{0}_report.docx'.format(alpha_id))
